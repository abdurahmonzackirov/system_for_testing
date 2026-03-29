"""from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

document = Document()

title = document.add_paragraph()
run1 = title.add_run('Dear Kate and Nick,')
font = run1.font
font.name = 'Arial'
font.size = Pt(11)
run1.italic = True
title_format = title.paragraph_format
title_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

body1 = document.add_paragraph()
continue1 = body1.add_run('We are looking forward very much to \
your visit to our country this summer. We are expecting you at the \
beginning of July and are hoping \
that you may stay until the end of the month or longer.')
font1 = continue1.font
font1.name = 'Times New Roman'
font1.size = Pt(12)
body1_format = body1.paragraph_format
body1_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
body1_format.first_line_indent = Inches(0.5)

body2 = document.add_paragraph()
continue2 = body2.add_run('We consider it a privilege for us to receive you as guests in our house. We are very grateful indeed to you for consenting to come and stay with us. We are looking forward to offering you hospitality in return for the hospitality you have kindly given us on many occasions.')
font2 = continue2.font
font2.name = 'Times New Roman'
font2.size = Pt(12)
body2_format = body2.paragraph_format
body2_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
body2_format.first_line_indent = Inches(0.5)

body3 = document.add_paragraph('')
continue3 = body3.add_run('We want you to understand that we will see to all your needs while you are with us and to any expenses that may arise.')
font3 = continue3.font
font3.name = 'Times New Roman'
font3.size = Pt(12)
body3_format = body3.paragraph_format
body3_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
body3_format.first_line_indent = Inches(0.5)

closing = document.add_paragraph('Yours sincerely,')
closing_format = closing.paragraph_format
closing_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
signature = closing.add_run('\nJohn and Mary.')
signature.bold = True

document.save('letter.docx')"""


class Summator:
    def __init__(self, n, N):
        self.n = n
        self.N = N
    
    def transform(self, operation=lambda x: x):
        return [operation(i) for i in range(self.n, self.N + 1)]
    
    def sum(self, operation=lambda x: x):
        return sum(self.transform(operation))
    

class SquareSummator(Summator):
    def __init__(self, n, N):
        super().__init__(n, N)
    
    def sum(self):
        return super().sum(lambda x: x ** 2)
    

class CubeSummator(Summator):
    def __init__(self, n, N):
        super().__init__(n, N)
    
    def sum(self):
        return super().sum(lambda x: x ** 3)
    

if __name__ == "__main__":
    n = 1
    N = 5
    summator = Summator(n, N)
    print(f"Sum of numbers from {n} to {N}: {summator.sum()}")
    
    square_summator = SquareSummator(n, N)
    print(f"Sum of squares from {n} to {N}: {square_summator.sum()}")
    
    cube_summator = CubeSummator(n, N)
    print(f"Sum of cubes from {n} to {N}: {cube_summator.sum()}")